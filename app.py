import streamlit as st
import os
from uuid import uuid4
from docx import Document as DocxDocument
import fitz  # PyMuPDF for PDF reading
import google.generativeai as genai
from pydantic import BaseModel, Field
from langchain_core.documents import Document
from pinecone import Pinecone, ServerlessSpec, PineconeApiException
import time
from langchain_pinecone import PineconeVectorStore
from langchain_google_genai import GoogleGenerativeAIEmbeddings

# Initialize Google GenAI API (API key will be provided in the sidebar)
# This model is used for generating responses and analyzing candidate data
model = genai.GenerativeModel(model_name="gemini-1.5-flash")

# Define schemas for Google GenAI response based on the type of content being analyzed
# Each class defines expected fields for structured responses from GenAI API

# Schema for resume analysis, covering key attributes in the resume
class CV(BaseModel):
    non_technical_skills: str = Field(description="Non-technical skills of the person")
    personal_details: str = Field(description="Personal details of the person")
    technical_skills: str = Field(description="Technical skills of the person")
    achievement: str = Field(description="Achievements of the person")
    certifications: str = Field(description="Certifications obtained by the person")
    qualifications: str = Field(description="Educational qualifications of the person")
    projects: str = Field(description="Projects undertaken by the person")
    past_experiences: str = Field(description="Previous work experiences of the person")
    quality_of_past_experiences: str = Field(description="Assessment of the quality of past experiences")
    extra_curricular_activities: str = Field(description="Extra-curricular activities participated in by the person")

# Schema for video analysis, focusing on attributes that can be derived from video content
class Video(BaseModel):
    life_experience: str = Field(description="Life experience of the person")
    name: str = Field(description="Name of the person")
    address: str = Field(description="Address of the person")
    qualification: str = Field(description="Educational qualification of the person")
    certification: str = Field(description="Certifications achieved by the person")
    experience: str = Field(description="Work experience of the person")
    confidence_tone: int = Field(description="Tone indicating the confidence level of the person 1 being highest 10 being least")
    vocabulary: str = Field(description="Vocabulary usage in the document")
    expressibility: str = Field(description="Ability to express ideas clearly")
    goal: str = Field(description="Goals of the person")
    aim: str = Field(description="Aims of the person")
    achievement: str = Field(description="Achievements of the person")
    passion: str = Field(description="Passions or interests of the person")
    job_seeking_for: str = Field(description="The type of job the person is seeking")
    non_technical_skills: str = Field(description="Non-technical skills of the person")
    technical_skills: str = Field(description="Technical skills of the person")
    personal_details: str = Field(description="Personal details of the person")

# Schema for final evaluation analysis, which provides recommendations for improving the candidate’s profile
class Evaluate(BaseModel):
    general_comments: str = Field(description="General commentary on how well it matches the given JD")
    areas_to_improve: str = Field(description="Suggestions on areas to improve.")


# Delete function to remove specific vectors from the Pinecone vector store
def delete(vector_store, uuids):
    for i in range(len(uuids)):
        vector_store.delete(ids=[uuids[i]])

# Document reading functions for handling DOCX and PDF formats
# Reads DOCX files, returning the text content as a single string
def read_docx(file_path):
    doc = DocxDocument(file_path)
    content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return content

# Reads PDF files using PyMuPDF, returning text content as a single string
def read_pdf(file_path):
    content = ""
    with fitz.open(file_path) as pdf:
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            content += page.get_text()
    return content

# Function to process resume content using Google GenAI, expecting structured output
def process_resume(content):
    result = model.generate_content(
        [content, '''Analyze this resume with a focus on skills, qualifications, achievements, and work experience. Provide a 
        structured summary that includes non-technical skills, technical skills, certifications, qualifications, projects, 
        past experiences, and extra-curricular activities. Ensure clarity and precision in each area.'''
],
        generation_config=genai.GenerationConfig(
            response_mime_type="application/json", response_schema=CV
        ), request_options={"timeout": 600}
    )
    return eval(result.text)

# Function to process video content using Google GenAI, uploading and analyzing it for key attributes
def process_video(video_file_name):
    video_file = genai.upload_file(path=video_file_name)
    while video_file.state.name == "PROCESSING":
        st.write("Processing video...")
        time.sleep(10)
        video_file = genai.get_file(video_file.name)
    result = model.generate_content(
        [video_file, '''Analyze this video for candidate attributes, focusing on life experience, educational background,
        certifications, work experience, personal qualities (such as confidence and expressibility), and career goals. 
        Provide a structured output with observations in each category.'''
        ],
        generation_config=genai.GenerationConfig(
            response_mime_type="application/json", response_schema=Video
        ), request_options={"timeout": 600}
    )
    return eval(result.text)

# Function to evaluate JD-Resume-Video alignment, outputting strengths and improvement areas
def send_to_genai_2(prompt):
    response = model.generate_content(
        [prompt, '''As an expert evaluator, analyze the alignment between the provided job description (JD), resume, and video.
        Identify key matching points, strengths, and areas for improvement in the candidate's skills, qualifications, experience,
        and expressed goals. Provide recommendations based on the match level.'''
        ],
        generation_config=genai.GenerationConfig(response_mime_type="application/json", response_schema=Evaluate),
        request_options={"timeout": 600}
    )
    return response.text

# Set up Streamlit page and sidebar configuration for API keys and file uploads
st.set_page_config(page_title="Job Matching and Interview Analysis")
st.sidebar.header("Settings")

# Sidebar API keys
st.sidebar.subheader("API Keys")
google_api_key = st.sidebar.text_input("Google GenAI API Key", type="password")
pinecone_api_key = st.sidebar.text_input("Pinecone API Key", type="password")


# Sidebar uploads for job descriptions, resume, and video files
st.sidebar.header("File Uploads")
job_files = st.sidebar.file_uploader("Upload Job Descriptions (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True)
resume_file = st.sidebar.file_uploader("Upload Resume (PDF/DOCX)", type=["pdf", "docx"])
video_file = st.sidebar.file_uploader("Upload Video (MP4/AVI/MOV)", type=["mp4", "avi", "mov"])

# Display tabs for different sections: Match Results, Job Description, Resume Analysis, and Video Analysis
st.title("Job Matching and Candidate Analysis")
tab1, tab2, tab3, tab4 = st.tabs(["Match Results","Job Description", "Resume Analysis", "Video Analysis"])

# Display sections based on the selected tab
with tab1:
    st.subheader("Match Results")
    st.write("Match results will display here once job descriptions, resume, and video are processed.")

with tab2:
    st.subheader("Job Descriptions")
    if job_files:
        for job_file in job_files:
            st.write(f"Uploaded: {job_file.name}")
    else:
        st.info("Please upload job descriptions.")

with tab3:
    st.subheader("Resume Analysis")
    if resume_file:
        st.write(f"Uploaded: {resume_file.name}")
    else:
        st.info("Please upload a resume.")

with tab4:
    st.subheader("Video Analysis")
    if video_file:
        st.write(f"Uploaded: {video_file.name}")
    else:
        st.info("Please upload a video.")


# Check if both API keys are provided before running the main logic
if google_api_key and pinecone_api_key:
    genai.configure(api_key=google_api_key)
    pc = Pinecone(api_key=pinecone_api_key)
    index_name = "job-opening-index"
    try:
        # Initialize Pinecone index and embeddings for the job matching analysis
        index = pc.Index(index_name) 
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=google_api_key)
        vector_store = PineconeVectorStore(index=index, embedding=embeddings)
        
        # Upload and analyze files
        if st.sidebar.button("Analyze"):
            documents = []
            for job_file in job_files:
                temp_path = f"/tmp/{job_file.name}"
                with open(temp_path, "wb") as f:
                    f.write(job_file.getbuffer())
                content = read_pdf(temp_path) if job_file.type == "application/pdf" else read_docx(temp_path)
                documents.append(Document(page_content=content, metadata={"source": job_file.name}))

            # Add documents to Pinecone
            uuids = [str(uuid4()) for _ in documents]
            vector_store.add_documents(documents=documents, ids=uuids)
            st.success("Job descriptions added to the database.")

            # Delete documents from Pinecone index
            if st.sidebar.button("Delete"):
                pc.delete_index(index_name)
                st.sidebar.write("Delete button pressed")
                delete(vector_store, uuids)
                st.success("Job descriptions deleted from the database.")

            # Process resume and video if uploaded
            resume_analysis, video_analysis = None, None
            if resume_file:
                resume_content = read_pdf(resume_file) if resume_file.type == "application/pdf" else read_docx(resume_file)
                resume_analysis = process_resume(resume_content)
                
            if video_file:
                temp_video_path = f"/tmp/{video_file.name}"
                with open(temp_video_path, "wb") as f:
                    f.write(video_file.getbuffer())
                video_analysis = process_video(temp_video_path)

            # Retrieve match results and display structured output for analysis
            if resume_analysis and video_analysis:
                results = vector_store.similarity_search_with_score(str(resume_analysis) + str(video_analysis), k=1)
                res, score = results[0]
                matched_role = res.metadata["source"]
                matched_JD = res.page_content
                response = send_to_genai_2(f"JD: {matched_JD}\nRole: {matched_role}\nResume: {resume_analysis}\nVideo: {video_analysis}")

                # Display structured output for match results
                st.write("**Job Role Matched:**", matched_role)
                st.write("**Match Score:**", f"{score:.2%}")

                # Display analysis sections for resume and video content
                st.write("### Analysis")
                for section, content in eval(response).items():
                    st.write(f"**{section.replace('_', ' ').title()}:**")
                    st.markdown(f"> {content}")
                    
    except PineconeApiException as e:
        st.error(f"An error occurred: {str(e)}")
else:
    st.warning("Please enter both API keys to activate the interface.")



# Helper functions for reading document content in DOCX and PDF formats
def read_docx(file_path):
    doc = DocxDocument(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

def read_pdf(file_path):
    content = ""
    with fitz.open(file_path) as pdf:
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            content += page.get_text()
    return content
